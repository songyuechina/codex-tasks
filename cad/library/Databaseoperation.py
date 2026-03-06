# -*- coding: utf-8 -*-
#D:/codex-tasks/cad/library/Databaseoperation.py
#版本V2.0
import os
from openai import OpenAI

import subprocess

import shutil

import time

import mysql.connector

import csv
import re

import regex as re

from docx import Document

import win32com.client

import psutil

from typing import List, Tuple

import pandas as pd

from sqlalchemy import create_engine

import pymysql

# ==============================================================================
#            MYSQL DATABASE API MANIFEST / 数据库通用函数签名清单
# ==============================================================================
"""
1. 数据库生命周期管理 (Database Lifecycle & Management)
--------------------------------------------------------------------------------
[DB-MGMT-001] connect_to_db_no_db()
              : [连接] 连接数据库。

[DB-MGMT-002] connect_to_db(database)
              : [连接] 连接到指定数据库。
[DB-MGMT-003] ensure_connection_alive(conn, database=None, max_retry=3, retry_wait=1.0)
              : [连接] 确保数据库连接可用。
              
[DB-MGMT-004] create_database_if_not_exists(database_name)
              : [初始化] 如果数据库不存在则创建，默认字符集 UTF8。
              
[DB-MGMT-005] delete_database(database)
              : [危险] 彻底删除指定数据库 (DROP DATABASE)。
              
[DB-MGMT-006] export_database(database, output_file)
              : [备份] 调用 mysqldump 将整个数据库导出为 .sql 文件。
              : output_file: 绝对路径 (如 "C:/Backups/db.sql")。
              
[DB-MGMT-007] import_database(database, input_file)
              : [恢复] 调用 mysql 命令行将 .sql 文件导入指定数据库。
              
[DB-MGMT-008] login_and_check_database(database)
              : [检查] 连接数据库，列出所有表，并显示部分数据以验证连接。

2. 表结构定义与管理 (Table Structure DDL)
--------------------------------------------------------------------------------
[TBL-STRUCT-001] create_table(database, table_name, fields_with_comments)
                 : [通用建表] 创建带注释的表。
                 : fields_with_comments: 字典格式 {字段名: (类型, 注释)}
                 : 示例: {"name": ("VARCHAR(50)", "姓名")}
                 
[TBL-STRUCT-002] create_table_with_foreign_keys(database, table_name, fields, foreign_keys=None)
                 : [高级建表] 支持外键约束和级联删除 (ON DELETE CASCADE)。
                 : foreign_keys: [("本表字段", "外表名", "外表字段", "CASCADE")]
                 
[TBL-STRUCT-003] rename_table(database, old_table_name, new_table_name)
                 : [重命名] 修改表名。
                 
[TBL-STRUCT-004] add_column_to_table(database, table_name, column_name, column_type, comment="")
                 : [加列] 向现有表添加新字段。
                 
[TBL-STRUCT-005] rename_column(database, table_name, old_col, new_col, col_type)
                 : [改列] 重命名现有字段 (MySQL要求必须同时指定类型)。
                 
[TBL-STRUCT-006] reset_auto_increment(database, table_name, new_value)
                 : [维护] 重置自增主键的起始值。
                 
[TBL-STRUCT-007] chakan_biao_biaojiegou(database)
                 : [内省] 打印数据库中所有表的 DDL (Create Statement)。

3. 数据写入引擎 (Data Insertion - Create)
--------------------------------------------------------------------------------
[DATA-INS-001] insert_record(database, table_name, data: dict)
               : [单条] 插入单条记录。data 为字典 {"字段": 值}。
               
[DATA-INS-002] insert_records_batch(database, table_name, records: list[dict])
               : [批量-推荐] 批量插入字典列表。效率优于单条插入。
               
[DATA-INS-003] append_records(database, table_name, records: list[tuple])
               : [追加-元组] 批量追加元组数据 (自动跳过ID列，依赖列顺序)。
               
[DATA-INS-004] append_records_with_id(database, table_name, records: list[tuple])
               : [追加-含ID] 批量追加包含ID的元组数据 (用于数据迁移)。
               
[DATA-INS-005] json_to_mysql_batch_insert(json_path, database, table_name)
               : [源-JSON] 从 JSON 文件 (List[Dict]) 导入数据。
               
[DATA-INS-006] excel_to_mysql_batch_insert(excel_path, database, table_name, user="root", host="localhost")
               : [源-Excel] 从 Excel 文件追加数据 (依赖 Pandas)。

4. 数据查询引擎 (Data Retrieval - Read)
--------------------------------------------------------------------------------
[DATA-READ-001] show_records_id(database, table_name, record_id=None)
                : [查全/查单] 不传ID显示全表(慎用)，传ID显示单条。返回记录列表。
                
[DATA-READ-002] show_records_range(database, table_name, start, end)
                : [分页] 查询 ID 范围内的记录 (LIMIT/OFFSET)。
                
[DATA-READ-003] show_records_by_field(database, table_name, field_name, field_value)
                : [精确] 根据指定字段值查询 (WHERE field = value)。
                
[DATA-READ-004] batch_show_records_by_field(database, table_name, field_name, values_list)
                : [批量-精确] 循环查询列表中的每个值。
                
[DATA-READ-005] find_words_containing_substring(database, substring)
                : [模糊] 搜索包含指定子串的记录 (LIKE %str%)。
                
[DATA-READ-006] get_related_records(database, base_table, base_table_id)
                : [关联] 查询所有引用了 base_table_id 的子表记录。

5. 数据修改与删除 (Update & Delete)
--------------------------------------------------------------------------------
[DATA-MOD-001] update_field_info(database, table_name, record_id, field, new_value)
               : [更新-单值] 修改指定 ID 的单个字段。
               
[DATA-MOD-002] batch_update_records(database, table_name, fields, *values_lists)
               : [更新-批量] 基于第一列作为 Key 批量更新后续列。
               
[DATA-MOD-003] delete_record_by_id(database, table, id_value)
               : [删除-ID] 删除指定主键的记录。
               
[DATA-MOD-004] delete_records(database, table, where_field, values, operator="IN")
               : [删除-批量] 通用批量删除 (如删除 room_id IN [1,2,3])。
               
[DATA-MOD-005] delete_records_with_conditions(database, table, conditions: dict)
               : [删除-多条件] AND 逻辑删除 (如 {"type": "A", "status": 0})。
               
[DATA-MOD-006] delete_records_advanced(database, table, conditions)
               : [删除-高级] 支持复杂操作符 (BETWEEN, LIKE, !=)。
               : 示例: {"age": ("BETWEEN", (10, 20))}
               
[DATA-MOD-007] delete_last_n_records(database, table_name, n)
               : [清理] 删除表中最后插入的 N 条数据。

6. 数据迁移与转换 (Migration & Transfer)
--------------------------------------------------------------------------------
[DATA-XFER-001] export_mysql_to_excel(user, host, database, output_path)
                : [导出-全库] 将库中所有表导出为 Excel (多Sheet)。
                
[DATA-XFER-002] export_mysql_schema_to_excel(user, host, database, output_path)
                : [导出-结构] 仅导出数据库表结构定义到 Excel。
                
[DATA-XFER-003] restore_database_from_excel(excel_path, database, ...)
                : [恢复-数据] 从 Excel 恢复数据 (Sheet名=表名)。
                
[DATA-XFER-004] restore_database_from_schema_and_data(schema_path, data_path, database_name, ...)
                : [完全恢复] 基于结构 Excel 和数据 Excel 重建整个数据库。
"""

__all__ = [
    # 数据库管理
    'connect_to_db_no_db','connect_to_db','ensure_connection_alive','create_database_if_not_exists', 'delete_database', 'export_database', 'import_database', 'login_and_check_database',
    # 表结构
    'create_table', 'create_table_with_foreign_keys', 'rename_table', 'add_column_to_table', 'rename_column', 'reset_auto_increment', 'chakan_biao_biaojiegou',
    # 插入
    'insert_record', 'insert_records_batch', 'append_records', 'append_records_with_id', 'json_to_mysql_batch_insert', 'excel_to_mysql_batch_insert',
    # 查询
    'show_records_id', 'show_records_range', 'show_records_by_field', 'batch_show_records_by_field', 'find_words_containing_substring', 'get_related_records',
    # 修改删除
    'update_field_info', 'batch_update_records', 'delete_record_by_id', 'delete_records', 'delete_records_with_conditions', 'delete_records_advanced', 'delete_last_n_records',
    # 迁移
    'export_mysql_to_excel', 'export_mysql_schema_to_excel', 'restore_database_from_excel', 'restore_database_from_schema_and_data'
]

# 连接到 MySQL 服务器，不指定数据库
def connect_to_db_no_db():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password=os.getenv("DATABASE_PASSWORD")
    )

# 创建一个简单的未定义数据结构关系的数据库
def create_database_if_not_exists(database_name):
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")

    try:
        db = mysql.connector.connect(host=host, user=user, password=password)
        cursor = db.cursor()

        cursor.execute(f"SHOW DATABASES LIKE %s", (database_name,))
        result = cursor.fetchone()

        if result:
            print(f"Database {database_name} already exists.")
        else:
            cursor.execute(
                f"CREATE DATABASE `{database_name}` "
                f"DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci;"
            )
            print(f"Database {database_name} created successfully.")

        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

# 连接到指定的数据库
def connect_to_db(database):
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password=os.getenv("DATABASE_PASSWORD"),
        database=database
    )

# 确保数据库连接可用
def ensure_connection_alive(conn, database=None, max_retry=3, retry_wait=1.0):
    """
    确保 MySQL 连接可用。如果连接已断开，自动重连。

    参数:
        conn: mysql.connector connection 对象（可能已经失效）
        database: 要连接的数据库名（可选，如果为 None 则只连 server）
        max_retry: 最大重试次数
        retry_wait: 每次重试间隔（秒）

    返回:
        一个【保证可用】的 mysql 连接对象

    行为:
        - 如果 conn 还活着：直接返回 conn
        - 如果断了：尝试重连
        - 如果多次重连失败：抛异常
    """

    def _is_alive(c):
        try:
            # 这个是官方推荐的轻量心跳
            c.ping(reconnect=False, attempts=1, delay=0)
            return True
        except:
            return False

    # --- 1. 如果原连接存在且可用，直接返回 ---
    if conn is not None:
        try:
            if _is_alive(conn):
                return conn
        except:
            pass  # 继续走重连逻辑

    # --- 2. 需要重连 ---
    last_err = None

    for i in range(1, max_retry + 1):
        try:
            if database:
                new_conn = mysql.connector.connect(
                    host="localhost",
                    user="root",
                    password=os.getenv("DATABASE_PASSWORD"),
                    database=database,
                )
            else:
                new_conn = mysql.connector.connect(
                    host="localhost",
                    user="root",
                    password=os.getenv("DATABASE_PASSWORD"),
                )

            # 再确认一次
            new_conn.ping(reconnect=False, attempts=1, delay=0)

            # 成功
            print(f"[DB] Reconnected successfully (attempt {i})")
            return new_conn

        except Exception as e:
            last_err = e
            print(f"[DB] Reconnect failed ({i}/{max_retry}): {e}")
            time.sleep(retry_wait)

    # --- 3. 彻底失败 ---
    raise RuntimeError(f"Cannot reconnect to MySQL after {max_retry} attempts: {last_err}")

def execute_sql(database: str, sql: str, params=None, *, fetch: bool=False, dictionary: bool=False):
    db = connect_to_db(database)
    cur = db.cursor(dictionary=dictionary)
    try:
        cur.execute(sql, params or ())
        if fetch:
            return cur.fetchall()
        db.commit()
        return True
    finally:
        cur.close()
        db.close()



# 在指定的数据库中创建 WordBasicInfo 表
def create_table(database):
    db = connect_to_db(database)
    cursor = db.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS WordBasicInfo (
            id INT AUTO_INCREMENT PRIMARY KEY,
            code INT,
            word VARCHAR(255),
            phonetic VARCHAR(255),
            meaning TEXT,
            UNIQUE (code)
        )
    """)
    cursor.close()
    db.close()





def insert_word_basic_info(database, code, word, phonetic, meaning):
    db = connect_to_db(database)
    cursor = db.cursor()
    sql = "INSERT INTO WordBasicInfo (code, word, phonetic, meaning) VALUES (%s, %s, %s, %s)"
    cursor.execute(sql, (code, word, phonetic, meaning))
    db.commit()
    cursor.close()
    db.close()

def parse_line(line):
    # 调整正则表达式以匹配check_document_format中的模式
    pattern = re.compile(r"(\d+)\.([a-zA-Z]+(\.[a-zA-Z]*)*(-[a-zA-Z]*)*)\s?\[([^\u4e00-\u9fa5\[\]]+)\]([a-zA-Z]+)\.\s?(.*)")
    pattern_with_apostrophe = re.compile(r"(\d+)\.([a-zA-Z\']+(\.[a-zA-Z]*)*(-[a-zA-Z]*)*)\s?\[([^\u4e00-\u9fa5\[\]]+)\]([a-zA-Z]+)\.\s?(.*)")
    specific_pattern = re.compile(r"(\d+)\.([a-zA-Z\.]+)\[([^\u4e00-\u9fa5\[\]]+)\]([a-zA-Z]+)\.\s?(.*)")

    match = pattern.match(line) or pattern_with_apostrophe.match(line) or specific_pattern.match(line)
    if match:
        code = int(match.group(1))
        word = match.group(2)
        phonetic = match.group(5)
        meaning = match.group(6) + "." + match.group(7)
        return code, word, phonetic, meaning
    return None

def process_document(database, doc_path):
    create_table(database)
    doc = Document(doc_path)
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parsed_data = parse_line(line)
            if parsed_data:
                insert_word_basic_info(database, *parsed_data)
                print(f"Inserted: {parsed_data}")
            else:
                print(f"Failed to parse line: {line}")

# 删除指定的数据库
def delete_database(database):
    db = connect_to_db_no_db()
    cursor = db.cursor()
    cursor.execute(f"DROP DATABASE IF EXISTS {database}")
    db.commit()
    cursor.close()
    db.close()

##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


### 二            数据库的创建、连接、查看


##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&






## 2 创建English_wordsstudy复杂表结构  别的数据库参考这个操作
def create_tables_fuza(database):
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")

    try:
        # 连接到数据库
        db = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()
        
        # 创建表
        table_creation_statements = [
            """
            CREATE TABLE IF NOT EXISTS Projects (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(255) NOT NULL,
                description TEXT
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS Clusters (
                id INT AUTO_INCREMENT PRIMARY KEY,
                project_id INT,
                name VARCHAR(255) NOT NULL,
                description TEXT,
                FOREIGN KEY (project_id) REFERENCES Projects(id)
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS `Groups` (
                id INT AUTO_INCREMENT PRIMARY KEY,
                cluster_id INT,
                name VARCHAR(255) NOT NULL,
                description TEXT,
                FOREIGN KEY (cluster_id) REFERENCES Clusters(id)
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS Words (
                id INT AUTO_INCREMENT PRIMARY KEY,
                group_id INT,
                word VARCHAR(255) NOT NULL,
                phonetic VARCHAR(255),
                meaning TEXT,
                FOREIGN KEY (group_id) REFERENCES `Groups`(id)
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS Details (
                id INT AUTO_INCREMENT PRIMARY KEY,
                word_id INT,
                part_of_speech VARCHAR(50),
                definition TEXT,
                example_sentences TEXT,
                analysis TEXT,
                FOREIGN KEY (word_id) REFERENCES Words(id)
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS Sentences (
                id INT AUTO_INCREMENT PRIMARY KEY,
                word_id INT,
                sentence TEXT,
                FOREIGN KEY (word_id) REFERENCES Words(id)
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS Images (
                id INT AUTO_INCREMENT PRIMARY KEY,
                word_id INT,
                image_path VARCHAR(255) NOT NULL,
                description TEXT,
                FOREIGN KEY (word_id) REFERENCES Words(id)
            );
            """
        ]

        for statement in table_creation_statements:
            try:
                cursor.execute(statement)
                print(f"Executed: {statement}")
            except mysql.connector.Error as err:
                print(f"Error executing: {statement}")
                print(f"Error: {err}")
        
        db.commit()
        print(f"Tables created successfully in {database}.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

## 3 数据库的登陆和基本查询

def login_and_check_database(database):
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        # 连接到MySQL服务器
        db = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()
        
        # 切换到新的数据库
        cursor.execute(f"USE {database}")
        print(f"Using database: {database}")
        
        # 列出数据库中的所有表
        cursor.execute("SHOW TABLES")
        tables = cursor.fetchall()
        print("Tables in the database:")
        for table in tables:
            print(table[0])
        
        # 检查WordBasicInfo表中的数据
        cursor.execute("SELECT COUNT(*) FROM WordBasicInfo")
        count = cursor.fetchone()[0]
        print(f"Total records in WordBasicInfo: {count}")
        
        cursor.execute("SELECT * FROM WordBasicInfo LIMIT 10")
        rows = cursor.fetchall()
        print("First 10 records in WordBasicInfo:")
        for row in rows:
            print(row)
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")








## 查看表和表结构

def chakan_biao_biaojiegou(database):
    host     = "localhost"
    user     = "root"
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        db     = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()
        
        # 1) 取所有表名
        cursor.execute("SHOW TABLES")
        tables = [t[0] for t in cursor.fetchall()]
        
        # 2) 逐表打印 DDL
        for table in tables:
            cursor.execute(f"SHOW CREATE TABLE `{table}`")
            ddl = cursor.fetchone()[1]
            
            print(f"— 表 `{table}` DDL —")
            # 方法 A：直接 print，内部的 \n 会被当成换行
            print(ddl)
            
            # 如果想更精细地逐行打印，也可以用下面这段：
            # for line in ddl.splitlines():
            #     print(line)
            
            print()  # 空行分隔
        
        cursor.close()
        db.close()
        
    except mysql.connector.Error as err:
        print(f"Error: {err}")


def show_create_table(database: str, table_name: str) -> str:
    """
    返回指定表的 CREATE TABLE DDL（等价于 SHOW CREATE TABLE）。
    """
    db = connect_to_db(database)
    cursor = db.cursor()
    try:
        cursor.execute(f"SHOW CREATE TABLE `{table_name}`")
        row = cursor.fetchone()
        if not row:
            return ""
        # row[0]=table, row[1]=ddl
        return row[1]
    finally:
        cursor.close()
        db.close()


def show_columns(database: str, table_name: str) -> list[tuple]:
    """
    返回指定表的字段信息（等价于 SHOW COLUMNS）。
    """
    db = connect_to_db(database)
    cursor = db.cursor()
    try:
        cursor.execute(f"SHOW COLUMNS FROM `{table_name}`")
        return cursor.fetchall()
    finally:
        cursor.close()
        db.close()
        

#查看重复键

# ==============================================================================
# FUNCINFO 专用：检查 (qualified_name, source_hash) 是否重复
# 等价于：
# SELECT qualified_name, source_hash, COUNT(*) AS c
# FROM function_analysis
# GROUP BY qualified_name, source_hash
# HAVING c > 1;
# ==============================================================================

def check_function_analysis_duplicates(database="CAD_FUNCINFO", verbose=True):
    """
    检查 function_analysis 表中
    (qualified_name, source_hash) 是否存在重复。

    返回:
        {
            "ok": True/False,
            "duplicates": list[dict],
            "count": int
        }
    """

    conn = None
    cursor = None

    try:
        conn = connect_to_db(database)
        cursor = conn.cursor(dictionary=True)

        sql = """
        SELECT
            qualified_name,
            source_hash,
            COUNT(*) AS c
        FROM function_analysis
        GROUP BY qualified_name, source_hash
        HAVING c > 1;
        """

        cursor.execute(sql)
        rows = cursor.fetchall()

        if verbose:
            print("\n" + "="*80)
            print("🔍 DUPLICATE CHECK — function_analysis")
            print("="*80)

            if not rows:
                print("✅ 无重复记录 (qualified_name, source_hash)")
            else:
                for r in rows:
                    print(
                        f"{r['qualified_name']}  |  "
                        f"{r['source_hash'][:12]}...  |  "
                        f"count={r['c']}"
                    )

        return {
            "ok": True,
            "duplicates": rows,
            "count": len(rows)
        }

    except Exception as e:
        print(f"❌ Duplicate check failed: {e}")
        return {
            "ok": False,
            "duplicates": [],
            "count": 0,
            "error": str(e)
        }

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


# ==============================================================================
# FUNCINFO 专用：删除重复，仅保留最新 id
# ==============================================================================

def deduplicate_function_analysis(
    database="CAD_FUNCINFO",
    dry_run=True,
    verbose=True
):
    """
    删除重复记录：
    同 (qualified_name, source_hash) 仅保留 id 最大的一条。

    参数:
        dry_run=True  → 只预览，不删除
        dry_run=False → 执行删除

    返回:
        删除条数统计
    """

    conn = None
    cursor = None

    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()

        # 找重复主键 id
        sql_find = """
        SELECT id
        FROM function_analysis
        WHERE (qualified_name, source_hash, id) NOT IN (
            SELECT
                qualified_name,
                source_hash,
                MAX(id)
            FROM function_analysis
            GROUP BY qualified_name, source_hash
        );
        """

        cursor.execute(sql_find)
        ids = [r[0] for r in cursor.fetchall()]

        if verbose:
            print("\n" + "="*80)
            print("🧹 DEDUP PREVIEW — function_analysis")
            print("="*80)
            print(f"Will delete rows: {len(ids)}")

        if dry_run:
            return {
                "ok": True,
                "to_delete": len(ids),
                "deleted": 0,
                "dry_run": True
            }

        # 真删除
        if ids:
            placeholders = ",".join(["%s"] * len(ids))
            sql_del = f"""
            DELETE FROM function_analysis
            WHERE id IN ({placeholders})
            """
            cursor.execute(sql_del, ids)
            conn.commit()

        return {
            "ok": True,
            "to_delete": len(ids),
            "deleted": len(ids),
            "dry_run": False
        }

    except Exception as e:
        print(f"❌ Deduplicate failed: {e}")
        return {
            "ok": False,
            "error": str(e)
        }

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==============================================================================
# FUNCINFO 专用：添加 uq_qname_source 唯一索引
# ==============================================================================

def add_unique_index_qname_source(
    database="CAD_FUNCINFO",
    dry_run=False
):
    """
    添加唯一键：
    (qualified_name(191), source_hash)

    注意：
        必须先确保无重复
    """

    sql = """
    ALTER TABLE `function_analysis`
      ADD UNIQUE KEY `uq_qname_source`
      (`qualified_name`(191), `source_hash`);
    """

    if dry_run:
        print(sql)
        return {"ok": True, "dry_run": True}

    conn = None
    cursor = None

    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()

        cursor.execute(sql)
        conn.commit()

        print("✅ uq_qname_source added")

        return {"ok": True}

    except Exception as e:
        print(f"❌ Add index failed: {e}")
        return {"ok": False, "error": str(e)}

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


#加唯一索引
def add_unique_index_qname_source(database="CAD_FUNCINFO", table="function_analysis", dry_run=False):
    """
    添加唯一索引： (qualified_name(191), source_hash)
    这是方案2的核心，让 DB 逻辑从 hash-key 转为可读 key。
    """

    sql = f"""
    ALTER TABLE `{table}`
      ADD UNIQUE KEY `uq_qname_source` (`qualified_name`(191), `source_hash`);
    """

    if dry_run:
        print(sql)
        return {"ok": True, "dry_run": True, "sql": sql}

    conn = None
    cursor = None
    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()
        cursor.execute(sql)
        conn.commit()
        print("✅ Added unique index: uq_qname_source (qualified_name(191), source_hash)")
        return {"ok": True, "error": None}
    except Exception as e:
        # 若已存在索引，会报错：Duplicate key name
        print(f"❌ Add uq_qname_source failed: {e}")
        return {"ok": False, "error": str(e)}
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


#查看索引

def show_table_indexes(database="CAD_FUNCINFO", table="function_analysis"):
    """
    打印/返回表索引信息：SHOW INDEX FROM table
    """
    conn = None
    cursor = None
    try:
        conn = connect_to_db(database)
        cursor = conn.cursor(dictionary=True)
        cursor.execute(f"SHOW INDEX FROM `{table}`;")
        rows = cursor.fetchall()

        print("\n" + "="*80)
        print(f"📌 INDEXES — {database}.{table}")
        print("="*80)
        for r in rows:
            print(
                f"{r.get('Key_name'):<20} "
                f"Non_unique={r.get('Non_unique')}  "
                f"Seq={r.get('Seq_in_index')}  "
                f"Column={r.get('Column_name')}  "
                f"Sub_part={r.get('Sub_part')}"
            )

        return {"ok": True, "indexes": rows}
    except Exception as e:
        print(f"❌ show_table_indexes failed: {e}")
        return {"ok": False, "error": str(e), "indexes": []}
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


#删除旧索引
def drop_index(database="CAD_FUNCINFO", table="function_analysis", index_name="uq_qhash_source", dry_run=False):
    """
    删除指定索引（可选清理步骤）。
    建议：确认 function_analyzer 已不依赖旧索引后再执行。
    """
    sql = f"ALTER TABLE `{table}` DROP INDEX `{index_name}`;"

    if dry_run:
        print(sql)
        return {"ok": True, "dry_run": True, "sql": sql}

    conn = None
    cursor = None
    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()
        cursor.execute(sql)
        conn.commit()
        print(f"✅ Dropped index: {index_name}")
        return {"ok": True}
    except Exception as e:
        print(f"❌ drop_index failed: {e}")
        return {"ok": False, "error": str(e)}
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

            
            
##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


###            数据库数据迁移



##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


# 导入数据库
def import_database(database, input_file):
    try:
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            password=os.getenv("DATABASE_PASSWORD"),
            database=database
        )
        cursor = db.cursor()
        
        # 禁用外键检查
        cursor.execute("SET FOREIGN_KEY_CHECKS = 0;")
        
        # 执行导入命令
        command = [
            'mysql',
            '--host', 'localhost',
            '--user', 'root',
            f'--password={os.getenv("DATABASE_PASSWORD")}',
            '--default-character-set=utf8',
            database
        ]
        
        with open(input_file, 'r', encoding='utf-8') as file:
            subprocess.run(command, stdin=file, check=True)
        
        # 启用外键检查
        cursor.execute("SET FOREIGN_KEY_CHECKS = 1;")
        
        print(f"Database {database} imported successfully from {input_file}.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    except subprocess.CalledProcessError as e:
        print(f"Error executing command: {e}")



##database = 'English_wordsstudy'
##
##input_file = 'C:/Users/Administrator/english_words_backup_utf8.sql'    
##    
##create_database(new_database)    
##    
##import_database(new_database, input_file)




##导出数据库

def export_database(database, output_file):
    # 使用mysqldump导出数据库
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")

    try:
        command = [
            'mysqldump',
            '--host', host,
            '--user', user,
            f'--password={password}',
            '--default-character-set=utf8',
            database
        ]
        with open(output_file, 'w', encoding='utf-8') as outfile:
            subprocess.run(command, stdout=outfile, check=True)
        print(f"Database {database} exported successfully to {output_file}.")

        return 1

    except subprocess.CalledProcessError as e:

        print(f"An error occurred: {e}")

        return 0


##if __name__ == "__main__":
##    host = 'localhost'
##    user = 'root'
##    password = 
##    database = 'english_words'
##    output_file = 'C:/Users/Administrator/english_words_backup_utf8.sql'
##    
##    export_database(host, user, password, database, output_file)
##





##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


###            更新表结构



##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&

## 1 备份现有数据

def backup_table(database, table_name, output_dir):##可能要修改
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        # 连接到数据库
        db = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()

        # 查询所有数据
        cursor.execute(f"SELECT * FROM `{table_name}`")
        rows = cursor.fetchall()

        # 获取列名
        cursor.execute(f"SHOW COLUMNS FROM `{table_name}`")
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]

        # 创建输出文件路径
        output_file = os.path.join(output_dir, f"{table_name}_backup.csv")

        # 写入CSV文件
        with open(output_file, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(column_names)
            writer.writerows(rows)

        cursor.close()
        db.close()
        print(f"Backup of {table_name} completed successfully.")
    except mysql.connector.Error as err:
        print(f"Error: {err}")

### 示例用法
##if __name__ == "__main__":
##    database = "English_wordsstudy"
##    output_directory = "E:/MySQL/"
##    tables_to_backup = ["clusters", "details", "groups", "images", "projects", "sentences", "wordbasicinfo", "words"]
##    
##    for table in tables_to_backup:
##        backup_table(database, table, output_directory)

## 2 删除现有表
def drop_existing_tables(database):
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        # 连接到数据库
        db = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()
        
        # 删除现有表
        cursor.execute("""
        DROP TABLE IF EXISTS `clusters`, `details`, `groups`, `images`, `projects`, `sentences`, `wordbasicinfo`, `words`;
        """)
        
        db.commit()
        print(f"Existing tables in {database} dropped successfully.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")


##创建建表结构的通用函数
"""
✅ 常用字段类型说明（MySQL）

类型名	含义和使用说明
INT	整数类型（范围约 ±21亿），可加 AUTO_INCREMENT 表示自增主键
FLOAT	浮点数，用于存储小数，如面积、宽度、坐标等
VARCHAR(n)	可变长度字符串，n 是最大字符数，如 VARCHAR(50) 表示最多可存 50 个字符（中文或英文）
TEXT	长文本字段（最大约 64KB），适合备注、大段描述等
JSON	JSON 格式数据（MySQL 5.7+ 支持），适合存储嵌套结构、坐标数组等
BOOLEAN	布尔值，实质是 TINYINT(1)，0 表示 False，1 表示 True
DATE	仅包含日期（如 2024-01-01）
DATETIME	日期和时间（如 2024-01-01 13:45:00）
TIMESTAMP	带时区的时间戳，通常用于记录创建或更新时间

类型	适合使用场景	示例字段
TINYINT(1)	表示布尔值（True/False）	is_isolated, has_toilet
ENUM	表示固定选项（空间类型等）	category ENUM('卧室','厨房','卫生间',...)
DATETIME	时间记录，如插入/修改时间	created_at, updated_at
GEOMETRY / POINT	如果你未来要做几何空间查询	position POINT（可与 GIS 引擎结合）
DOUBLE	比 FLOAT 更高精度坐标	高精度坐标需求时使用
BLOB	二进制大对象，如存 CAD 二进制片段	geometry_blob 等
CHAR(n)	固定长度字符串（如编码、状态码）	status_code CHAR(2)


name VARCHAR(50)
category ENUM('卧室','厨房','客厅','阳台','卫生间','通道') COMMENT '房间类型'
is_isolated TINYINT(1) DEFAULT 0 COMMENT '是否孤立'
created_at DATETIME DEFAULT CURRENT_TIMESTAMP
2. 墙体表推荐字段

width FLOAT
height FLOAT
usage ENUM('内墙','外墙','矮墙','虚墙') COMMENT '墙体类型'
start_point JSON
end_point JSON
3. 门窗表推荐字段

direction ENUM('内开','外开','平移','推拉') COMMENT '开启方式'
material VARCHAR(50)
position JSON  -- 插入点坐标

类型	是否推荐
INT	✅ 必需
FLOAT	✅ 推荐
DOUBLE	✅ 更精确
VARCHAR	✅ 必需
TEXT	✅ 长文本
TINYINT(1)	✅ 布尔型
ENUM	✅ 空间/类型分类字段
JSON	✅ 坐标数据或嵌套信息
DATETIME	✅ 记录更新时间
GEOMETRY	⚠️ 如果将来与 GIS 或地图系统对接



"""

def create_table(database, table_name, fields_with_comments):
    """
    在指定数据库中创建一个表，可附加字段注释。

    参数:
        database: 要使用的数据库名
        table_name: 表名
        fields_with_comments: dict，键为字段名，值为 (数据类型, 注释)

    示例:
        create_table(
            database="mydb",
            table_name="rooms",
            fields_with_comments={
                "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
                "name": ("VARCHAR(50)", "房间名称"),
                "category": ("VARCHAR(20)", "房间类型")
            }
        )
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 构造字段定义
        field_defs = []
        for field, (datatype, comment) in fields_with_comments.items():
            field_defs.append(f"`{field}` {datatype} COMMENT '{comment}'")

        # 拼接SQL
        sql = f"""
        CREATE TABLE IF NOT EXISTS `{table_name}` (
            {", ".join(field_defs)}
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """

        cursor.execute(sql)
        db.commit()
        print(f"✅ 成功创建表 `{table_name}`")

        cursor.close()
        db.close()
    except Exception as e:
        print(f"❌ 创建表失败：{e}")


##fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "name": ("VARCHAR(50)", "房间名称"),
##    "pl_handle": ("VARCHAR(20)", "多段线句柄"),
##    "pl_coordinates": ("JSON", "多段线坐标点"),
##    "category": ("VARCHAR(50)", "房间类型"),
##    "centerline_area": ("FLOAT", "中轴面积"),
##    "usable_area": ("FLOAT", "实用面积"),
##    "width": ("FLOAT", "开间"),
##    "depth": ("FLOAT", "进深"),
##    "remark": ("TEXT", "备注")
##}
##create_table("某某住宅", "rooms", fields)
##✅ 成功创建表 `rooms`
##chakan_biao_biaojiegou(database)
##Using database: 某某住宅
##Tables in the database:
##rooms
##{'rooms': "CREATE TABLE `rooms` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `name` varchar(50) DEFAULT NULL COMMENT '房间名称',\n  `pl_handle` varchar(20) DEFAULT NULL COMMENT '多段线句柄',\n  `pl_coordinates` json DEFAULT NULL COMMENT '多段线坐标点',\n  `category` varchar(50) DEFAULT NULL COMMENT '房间类型',\n  `centerline_area` float DEFAULT NULL COMMENT '中轴面积',\n  `usable_area` float DEFAULT NULL COMMENT '实用面积',\n  `width` float DEFAULT NULL COMMENT '开间',\n  `depth` float DEFAULT NULL COMMENT '进深',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci"}
##fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_from": ("INT", "起始房间ID"),
##    "room_to": ("INT", "目标房间ID"),
##    "relation_type": ("VARCHAR(50)", "关系类型（相邻、共墙、包含等）"),
##    "common_wall_handle": ("VARCHAR(20)", "共享墙句柄"),
##    "remark": ("TEXT", "备注")
##}
##create_table("某某住宅", "room_relations", fields)
##✅ 成功创建表 `room_relations`
##chakan_biao_biaojiegou(database)
##Using database: 某某住宅
##Tables in the database:
##room_relations
##rooms
##{'room_relations': "CREATE TABLE `room_relations` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `room_from` int DEFAULT NULL COMMENT '起始房间ID',\n  `room_to` int DEFAULT NULL COMMENT '目标房间ID',\n  `relation_type` varchar(50) DEFAULT NULL COMMENT '关系类型（相邻、共墙、包含等）',\n  `common_wall_handle` varchar(20) DEFAULT NULL COMMENT '共享墙句柄',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci", 'rooms': "CREATE TABLE `rooms` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `name` varchar(50) DEFAULT NULL COMMENT '房间名称',\n  `pl_handle` varchar(20) DEFAULT NULL COMMENT '多段线句柄',\n  `pl_coordinates` json DEFAULT NULL COMMENT '多段线坐标点',\n  `category` varchar(50) DEFAULT NULL COMMENT '房间类型',\n  `centerline_area` float DEFAULT NULL COMMENT '中轴面积',\n  `usable_area` float DEFAULT NULL COMMENT '实用面积',\n  `width` float DEFAULT NULL COMMENT '开间',\n  `depth` float DEFAULT NULL COMMENT '进深',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci"}
##fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID"),
##    "handle": ("VARCHAR(20)", "墙体句柄"),
##    "start_point": ("JSON", "起点坐标"),
##    "end_point": ("JSON", "终点坐标"),
##    "left_width": ("FLOAT", "左侧宽度"),
##    "right_width": ("FLOAT", "右侧宽度"),
##    "total_width": ("FLOAT", "总宽度"),
##    "usage": ("VARCHAR(50)", "墙体用途"),
##    "height": ("FLOAT", "墙高"),
##    "material": ("VARCHAR(50)", "材料"),
##    "remark": ("TEXT", "备注")
##}
##create_table("某某住宅", "walls", fields)
##✅ 成功创建表 `walls`
##door_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID"),
##    "handle": ("VARCHAR(20)", "门对象句柄"),
##    "position": ("JSON", "插入点坐标 (x, y, z)"),
##    "width": ("FLOAT", "门宽度"),
##    "open_direction": ("VARCHAR(50)", "开启方向（内开/外开）"),
##    "door_type": ("VARCHAR(50)", "门类型"),
##    "remark": ("TEXT", "备注")
##}
##create_table("某某住宅", "doors", door_fields)
##✅ 成功创建表 `doors`
##window_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID"),
##    "handle": ("VARCHAR(20)", "窗对象句柄"),
##    "position": ("JSON", "插入点坐标 (x, y, z)"),
##    "width": ("FLOAT", "窗宽度"),
##    "window_type": ("VARCHAR(50)", "窗类型"),
##    "remark": ("TEXT", "备注")
##}
##create_table("某某住宅", "windows", window_fields)
##✅ 成功创建表 `windows`
##chakan_biao_biaojiegou(database)
##Using database: 某某住宅
##Tables in the database:
##doors
##room_relations
##rooms
##walls
##windows
##{'doors': "CREATE TABLE `doors` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `room_id` int DEFAULT NULL COMMENT '所属房间ID',\n  `handle` varchar(20) DEFAULT NULL COMMENT '门对象句柄',\n  `position` json DEFAULT NULL COMMENT '插入点坐标 (x, y, z)',\n  `width` float DEFAULT NULL COMMENT '门宽度',\n  `open_direction` varchar(50) DEFAULT NULL COMMENT '开启方向（内开/外开）',\n  `door_type` varchar(50) DEFAULT NULL COMMENT '门类型',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci", 'room_relations': "CREATE TABLE `room_relations` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `room_from` int DEFAULT NULL COMMENT '起始房间ID',\n  `room_to` int DEFAULT NULL COMMENT '目标房间ID',\n  `relation_type` varchar(50) DEFAULT NULL COMMENT '关系类型（相邻、共墙、包含等）',\n  `common_wall_handle` varchar(20) DEFAULT NULL COMMENT '共享墙句柄',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci", 'rooms': "CREATE TABLE `rooms` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `name` varchar(50) DEFAULT NULL COMMENT '房间名称',\n  `pl_handle` varchar(20) DEFAULT NULL COMMENT '多段线句柄',\n  `pl_coordinates` json DEFAULT NULL COMMENT '多段线坐标点',\n  `category` varchar(50) DEFAULT NULL COMMENT '房间类型',\n  `centerline_area` float DEFAULT NULL COMMENT '中轴面积',\n  `usable_area` float DEFAULT NULL COMMENT '实用面积',\n  `width` float DEFAULT NULL COMMENT '开间',\n  `depth` float DEFAULT NULL COMMENT '进深',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci", 'walls': "CREATE TABLE `walls` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `room_id` int DEFAULT NULL COMMENT '所属房间ID',\n  `handle` varchar(20) DEFAULT NULL COMMENT '墙体句柄',\n  `start_point` json DEFAULT NULL COMMENT '起点坐标',\n  `end_point` json DEFAULT NULL COMMENT '终点坐标',\n  `left_width` float DEFAULT NULL COMMENT '左侧宽度',\n  `right_width` float DEFAULT NULL COMMENT '右侧宽度',\n  `total_width` float DEFAULT NULL COMMENT '总宽度',\n  `usage` varchar(50) DEFAULT NULL COMMENT '墙体用途',\n  `height` float DEFAULT NULL COMMENT '墙高',\n  `material` varchar(50) DEFAULT NULL COMMENT '材料',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci", 'windows': "CREATE TABLE `windows` (\n  `id` int NOT NULL AUTO_INCREMENT COMMENT '主键',\n  `room_id` int DEFAULT NULL COMMENT '所属房间ID',\n  `handle` varchar(20) DEFAULT NULL COMMENT '窗对象句柄',\n  `position` json DEFAULT NULL COMMENT '插入点坐标 (x, y, z)',\n  `width` float DEFAULT NULL COMMENT '窗宽度',\n  `window_type` varchar(50) DEFAULT NULL COMMENT '窗类型',\n  `remark` text COMMENT '备注',\n  PRIMARY KEY (`id`)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci"}
##



##给任何数据库的指定表添加字段

def add_column_to_table(database, table_name, column_name, column_type, comment=""):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        sql = f"""
        ALTER TABLE `{table_name}`
        ADD COLUMN `{column_name}` {column_type} COMMENT '{comment}';
        """
        cursor.execute(sql)
        db.commit()
        print(f"✅ 成功添加字段 `{column_name}` 到表 `{table_name}`")
        
        cursor.close()
        db.close()
    except Exception as e:
        print(f"❌ 添加字段失败：{e}")


##add_column_to_table(
##    database="某某住宅",
##    table_name="rooms",
##    column_name="is_isolated",
##    column_type="BOOLEAN DEFAULT FALSE",
##    comment="是否为孤立房间"
##)
##✅ 成功添加字段 `is_isolated` 到表 `rooms`







##恢复数据

def restore_table(database, table_name, input_dir):
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        # 连接到数据库
        db = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = db.cursor()

        # 创建输入文件路径
        input_file = os.path.join(input_dir, f"{table_name}_backup.csv")

        # 读取CSV文件
        with open(input_file, mode='r', encoding='utf-8') as file:
            reader = csv.reader(file)
            columns = next(reader)  # 获取列名
            query = f"INSERT INTO `{table_name}` ({', '.join(columns)}) VALUES ({', '.join(['%s'] * len(columns))})"
            
            for data in reader:
                cursor.execute(query, data)
        
        db.commit()
        print(f"Restored data to {table_name} successfully.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

### 示例用法
##if __name__ == "__main__":
##    database = "English_wordsstudy"
##    input_directory = "E:/MySQL/"
##    tables_to_restore = ["clusters", "details", "groups", "images", "projects", "sentences", "wordbasicinfo", "words"]
##    
##    for table in tables_to_restore:
##        restore_table(database, table, input_directory)









###数据库数据迁移到另一台电脑
##
##cmd的密码先CTR+C到cmd边上右键操作要干净
##
##cmd导出
##mysqldump -u root -p english_words > C:\Users\Administrator\english_words_backup.sql
##
##在本地电脑上
##C:/Users/Administrator/english_words_backup.sql
##将此文件移入另一台电脑同样位置
##
##cmd导入
##mysql -u root -p
##
##mysql
##CREATE DATABASE english_words;
##USE english_words;
##SOURCE C:\path\to\english_words_backup.sql;
##

### 第二部分  在基本信息基础上提取信息分别存入各表


## 提取信息插入各表

def extract_and_insert_data(database):
    db = connect_to_db(database)
    cursor = db.cursor()

    # 从 WordBasicInfo 表中提取信息
    cursor.execute("SELECT id, word, phonetic, meaning FROM WordBasicInfo ORDER BY id ASC")
    records = cursor.fetchall()

    for record in records:
        word_id = record[0]
        word = record[1]
        phonetic = f"[{record[2]}]"
        definition = record[3]

        # 插入到 Words 表中
        cursor.execute("INSERT INTO Words (id, word) VALUES (%s, %s)", (word_id, word))

        # 插入到 Phonetics 表中
        cursor.execute("INSERT INTO Phonetics (word_id, phonetic) VALUES (%s, %s)", (word_id, phonetic))

        # 插入到 Definitions 表中
        cursor.execute("INSERT INTO Definitions (word_id, definition) VALUES (%s, %s)", (word_id, definition))

    db.commit()
    cursor.close()
    db.close()



##从 WordBasicInfo 表中提取word信息值写入wordforms等

def extract_and_insert_wordforms(database):


    """
    当WordBasicInfo 表时更新时也可以用此函数将更新的值写入wordforms

    """    
    db = connect_to_db(database)
    cursor = db.cursor()

    # 从 WordBasicInfo 表中提取信息
    cursor.execute("SELECT id, word FROM WordBasicInfo ORDER BY id ASC")
    records = cursor.fetchall()

    for record in records:
        word_id = record[0]
        word = record[1]

        # 检查 WordForms 表中是否已经存在相应的记录
        cursor.execute("SELECT COUNT(*) FROM WordForms WHERE word_id = %s", (word_id,))
        count = cursor.fetchone()[0]

        if count > 0:
            # 更新现有记录
            cursor.execute("UPDATE WordForms SET word = %s WHERE word_id = %s", (word, word_id))
        else:
            # 插入新记录
            cursor.execute("INSERT INTO WordForms (word_id, word) VALUES (%s, %s)", (word_id, word))

    db.commit()
    cursor.close()
    db.close()


def extract_and_insert_phonetics(database):
    db = connect_to_db(database)
    cursor = db.cursor()

    # 从 WordBasicInfo 表中提取信息
    cursor.execute("SELECT id, phonetic FROM WordBasicInfo ORDER BY id ASC")
    records = cursor.fetchall()

    for record in records:
        word_id = record[0]
        phonetic = record[1]
        # 如果 phonetic 字段不为空，则加上方括号
        if phonetic:
            phonetic = f"[{phonetic}]"

        # 检查 phonetics 表中是否已经存在相应的记录
        cursor.execute("SELECT COUNT(*) FROM phonetics WHERE word_id = %s", (word_id,))
        count = cursor.fetchone()[0]

        if count > 0:
            # 更新现有记录
            cursor.execute("UPDATE phonetics SET phonetic = %s WHERE word_id = %s", (phonetic, word_id))
        else:
            # 插入新记录
            cursor.execute("INSERT INTO phonetics (word_id, phonetic) VALUES (%s, %s)", (word_id, phonetic))

    db.commit()
    cursor.close()
    db.close()

def extract_and_insert_definitions(database):
    db = connect_to_db(database)
    cursor = db.cursor()

    # 从 WordBasicInfo 表中提取信息
    cursor.execute("SELECT id, meaning FROM WordBasicInfo ORDER BY id ASC")
    records = cursor.fetchall()

    for record in records:
        word_id = record[0]
        definition = record[1]

        # 检查 definitions 表中是否已经存在相应的记录
        cursor.execute("SELECT COUNT(*) FROM definitions WHERE word_id = %s", (word_id,))
        count = cursor.fetchone()[0]

        if count > 0:
            # 更新现有记录
            cursor.execute("UPDATE definitions SET definition = %s WHERE word_id = %s", (definition, word_id))
        else:
            # 插入新记录
            cursor.execute("INSERT INTO definitions (word_id, definition) VALUES (%s, %s)", (word_id, definition))

    db.commit()
    cursor.close()
    db.close()






# 重命名表的函数
def rename_table(database, old_table_name, new_table_name):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        cursor.execute(f"ALTER TABLE {old_table_name} RENAME TO {new_table_name};")
        db.commit()
        cursor.close()
        db.close()
        print(f"Table {old_table_name} has been renamed to {new_table_name} successfully.")
    except mysql.connector.Error as err:
        print(f"Error: {err}")

def rename_column(database, table_name, old_column_name, new_column_name, column_type):
    """
    重命名指定表中的列字段。

    参数:
        database: 数据库名
        table_name: 表名
        old_column_name: 原列名
        new_column_name: 新列名
        column_type: 列的数据类型（必须指定，MySQL 重命名字段时必须提供类型）
    
    示例:
        rename_column("某某住宅", "rooms", "depth", "depth_m", "FLOAT")
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        sql = f"""
        ALTER TABLE `{table_name}` 
        CHANGE `{old_column_name}` `{new_column_name}` {column_type};
        """

        cursor.execute(sql)
        db.commit()
        print(f"✅ 列 `{old_column_name}` 已成功重命名为 `{new_column_name}`，类型为 {column_type}")

        cursor.close()
        db.close()

    except mysql.connector.Error as err:
        print(f"❌ 重命名列失败：{err}")


##支持 ON DELETE CASCADE 的增强版建表函数，用于自动删除关联主表记录时，同时删除从表中相关记录（如删除房间时，自动删除其门窗）

def create_table_with_foreign_keys(database, table_name, fields_with_comments, foreign_keys=None):
    """

    是否定义外键	是否限制房间 ID 必须存在？	删除房间会自动删除门/窗？
    没有外键（默认）	❌ 否，随便写	❌ 否，容易数据断裂
    有外键，无 ON DELETE	✅ 是，强制检查	❌ 否，删除房间时报错
    有外键，加 ON DELETE CASCADE	✅ 是	✅ 是，房间删了自动删子项




    创建带字段注释和外键约束的表，支持 ON DELETE CASCADE。

    参数:
        database: 数据库名
        table_name: 表名
        fields_with_comments: 字典，键为字段名，值为 (字段类型, 注释)
        foreign_keys: 外键列表，例如 [("room_id", "rooms", "id", "CASCADE")]

    示例:
        create_table_with_foreign_keys(
            "某某住宅",
            "doors",
            {
                "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
                "room_id": ("INT", "所属房间ID"),
                ...
            },
            [("room_id", "rooms", "id", "CASCADE")]
        )
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        field_defs = []
        for field, (dtype, comment) in fields_with_comments.items():
            field_defs.append(f"`{field}` {dtype} COMMENT '{comment}'")

        # 外键处理
        if foreign_keys:
            for fk in foreign_keys:
                fk_field, ref_table, ref_field = fk[:3]
                on_delete = fk[3] if len(fk) > 3 else "RESTRICT"
                fk_def = f"FOREIGN KEY (`{fk_field}`) REFERENCES `{ref_table}`(`{ref_field}`) ON DELETE {on_delete}"
                field_defs.append(fk_def)

        sql = f"""
        CREATE TABLE IF NOT EXISTS `{table_name}` (
            {", ".join(field_defs)}
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
        cursor.execute(sql)
        db.commit()
        print(f"✅ 成功创建表 `{table_name}`（含外键和级联删除）")
        cursor.close()
        db.close()
    except Exception as e:
        print(f"❌ 创建表失败：{e}")


##door_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID（引用 rooms.id）"),
##    "handle": ("VARCHAR(20)", "门对象句柄"),
##    "position": ("JSON", "插入点坐标 (x, y, z)"),
##    "width": ("FLOAT", "门宽度"),
##    "open_direction": ("VARCHAR(50)", "开启方向（内开/外开）"),
##    "door_type": ("VARCHAR(50)", "门类型"),
##    "remark": ("TEXT", "备注")
##}
##
##foreign_keys = [
##    ("room_id", "rooms", "id", "CASCADE")  # 删除房间时自动删除相关门
##]
##
##create_table_with_foreign_keys("某某住宅", "doors", door_fields, foreign_keys)
##
##window_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID（引用 rooms.id）"),
##    "handle": ("VARCHAR(20)", "窗对象句柄"),
##    "position": ("JSON", "插入点坐标 (x, y, z)"),
##    "width": ("FLOAT", "窗宽度"),
##    "window_type": ("VARCHAR(50)", "窗类型"),
##    "remark": ("TEXT", "备注")
##}
##window_foreign_keys = [("room_id", "rooms", "id", "CASCADE")]
##create_table_with_foreign_keys("某某住宅", "windows", window_fields, window_foreign_keys)
##
##wall_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_id": ("INT", "所属房间ID（引用 rooms.id）"),
##    "handle": ("VARCHAR(20)", "墙体句柄"),
##    "start_point": ("JSON", "起点坐标"),
##    "end_point": ("JSON", "终点坐标"),
##    "left_width": ("FLOAT", "左侧宽度"),
##    "right_width": ("FLOAT", "右侧宽度"),
##    "total_width": ("FLOAT", "总宽度"),
##    "usage": ("VARCHAR(50)", "墙体用途"),
##    "height": ("FLOAT", "墙高"),
##    "material": ("VARCHAR(50)", "材料"),
##    "remark": ("TEXT", "备注")
##}
##wall_foreign_keys = [("room_id", "rooms", "id", "CASCADE")]
##create_table_with_foreign_keys("某某住宅", "walls", wall_fields, wall_foreign_keys)
##
##
##room_relations_fields = {
##    "id": ("INT AUTO_INCREMENT PRIMARY KEY", "主键"),
##    "room_from": ("INT", "起始房间ID（引用 rooms.id）"),
##    "room_to": ("INT", "目标房间ID（引用 rooms.id）"),
##    "relation_type": ("VARCHAR(50)", "关系类型（相邻、共墙、包含等）"),
##    "common_wall_handle": ("VARCHAR(20)", "共享墙句柄"),
##    "remark": ("TEXT", "备注")
##}
##room_relations_foreign_keys = [
##    ("room_from", "rooms", "id", "SET NULL"),
##    ("room_to", "rooms", "id", "SET NULL")
##]
##create_table_with_foreign_keys("某某住宅", "room_relations", room_relations_fields, room_relations_foreign_keys)



#  主函数
#  (1)
# 给一个表录入一条信息

# 该函数系列包括如下一些函数


def insert_record(database, table_name, data: dict):
    """
    向指定数据库和表中插入一条记录。

    参数：
        database: 数据库名称
        table_name: 表名
        data: 字典形式的数据，如 {"name": "房间1", "width": 3000}

    返回：
        True 成功，False 失败
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 字段和对应值
        fields = ", ".join(f"`{key}`" for key in data.keys())
        placeholders = ", ".join(["%s"] * len(data))
        values = list(data.values())

        sql = f"INSERT INTO `{table_name}` ({fields}) VALUES ({placeholders})"
        cursor.execute(sql, values)

        db.commit()
        print(f"✅ 插入成功到 `{table_name}`：{data}")
        cursor.close()
        db.close()
        return True
    except Exception as e:
        print(f"❌ 插入失败：{e}")
        return False



##insert_record("某某住宅", "rooms", {
##    "name": "房间1",
##    "pl_handle": "123AB",
##    "pl_coordinates": '[ [0,0,0], [0,3000,0], [4000,3000,0], [4000,0,0], [0,0,0] ]',
##    "category": "卧室",
##    "centerline_area": 12.5,
##    "usable_area": 11.8,
##    "width": 4000,
##    "depth": 3000,
##    "remark": "测试用卧室"
##})
##
##
##import json
##
##insert_record("某某住宅", "rooms", {
##    "name": "房间2",
##    "pl_handle": "456CD",
##    "pl_coordinates": json.dumps([[1,2,0], [3,4,0]]),
##    ...
##})


#  主函数
#  (1)
# 批量录入信息

# 该函数系列包括如下一些函数


def insert_records_batch(database, table_name, records: list[dict]):
    """
    批量向指定数据库的指定表插入多条记录。

    参数：
        database: 数据库名称
        table_name: 表名
        records: 多个字典组成的列表，每个字典是1条数据，字段必须一致。

    返回：
        True 成功，False 失败
    """
    if not records:
        print("⚠️ 空记录列表，未插入任何数据")
        return False

    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 获取字段名和预编译语句
        keys = records[0].keys()
        fields = ", ".join(f"`{key}`" for key in keys)
        placeholders = ", ".join(["%s"] * len(keys))
        sql = f"INSERT INTO `{table_name}` ({fields}) VALUES ({placeholders})"

        # 构造值列表
        values_list = [tuple(rec[key] for key in keys) for rec in records]

        # 执行批量插入
        cursor.executemany(sql, values_list)

        db.commit()
        print(f"✅ 成功批量插入 {len(records)} 条记录到 `{table_name}`")
        cursor.close()
        db.close()
        return True

    except Exception as e:
        print(f"❌ 批量插入失败：{e}")
        return False

##
##batch_data = [
##    {
##        "name": "房间A",
##        "pl_handle": "111A",
##        "pl_coordinates": '[ [0,0,0], [0,3000,0], [4000,3000,0], [4000,0,0], [0,0,0] ]',
##        "category": "卧室",
##        "centerline_area": 12.5,
##        "usable_area": 11.8,
##        "width": 4000,
##        "depth": 3000,
##        "remark": "示例1"
##    },
##    {
##        "name": "房间B",
##        "pl_handle": "222B",
##        "pl_coordinates": '[ [1000,0,0], [1000,3000,0], [5000,3000,0], [5000,0,0], [1000,0,0] ]',
##        "category": "阳台",
##        "centerline_area": 5.5,
##        "usable_area": 5.0,
##        "width": 4000,
##        "depth": 3000,
##        "remark": "示例2"
##    }
##]
##
##insert_records_batch("某某住宅", "rooms", batch_data)

#  主函数
#  (1)
# EXCEL批量插入

# 该函数系列包括如下一些函数

import pandas as pd
from sqlalchemy import create_engine

def excel_to_mysql_batch_insert(excel_path, database, table_name, user="root", host="localhost"):
    """
    从 Excel 文件批量插入数据到指定数据库表。

    参数:
        excel_path: Excel 文件路径
        database: 目标数据库名称
        table_name: 目标表名称
        user: MySQL 用户名（默认 root）
        host: MySQL 主机地址（默认 localhost）
    """
    try:
        password = os.getenv("DATABASE_PASSWORD")
        df = pd.read_excel(excel_path)

        engine = create_engine(f"mysql+pymysql://{user}:{password}@{host}/{database}?charset=utf8mb4")

        df.to_sql(table_name, con=engine, if_exists='append', index=False)
        print(f"✅ 成功将 Excel 文件数据插入到 {database}.{table_name} 表中，共 {len(df)} 行。")
    except Exception as e:
        print(f"❌ Excel 插入失败：{e}")


#  主函数
#  (1)
# JSON批量插入

# 该函数系列包括如下一些函数

import json
import mysql.connector

def json_to_mysql_batch_insert(json_path, database, table_name):
    """
    从 JSON 文件中批量插入数据到指定表。

    参数:
        json_path: JSON 文件路径，格式为 list[dict]
        database: 数据库名
        table_name: 表名
    """
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            records = json.load(f)

        if not isinstance(records, list):
            print("❌ JSON 格式错误：必须为列表对象。")
            return

        db = connect_to_db(database)
        cursor = db.cursor()

        for row in records:
            columns = ", ".join(f"`{col}`" for col in row)
            placeholders = ", ".join(["%s"] * len(row))
            values = tuple(row.values())

            sql = f"INSERT INTO `{table_name}` ({columns}) VALUES ({placeholders})"
            cursor.execute(sql, values)

        db.commit()
        print(f"✅ 成功插入 {len(records)} 条记录到表 `{table_name}`。")

        cursor.close()
        db.close()
    except Exception as e:
        print(f"❌ JSON 插入失败：{e}")


##0618

### 从100个单词的列表，创建表groups的一个记录，其name字段以该列表第一个单词写入，其id作为100个单词的外键


##一个列表LB_word不在LB1的元素
def find_missing_element(LB1, LB_word):##临时位置
    # 将LB1和LB_word中的元素转换为小写，并转换为集合
    set_LB1 = set(word.lower() for word in LB1)
    set_LB_word = set(word.lower() for word in LB_word)
    
    # 找出LB_word中有但LB1中没有的元素
    missing_element = list(set_LB_word - set_LB1)
    
    return missing_element



##查询word表中不区分大小写情况下重复的单词
def find_duplicate_words(database):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 查询所有单词，转换为小写并统计出现次数
        cursor.execute("SELECT LOWER(word) FROM words")
        words = cursor.fetchall()
        
        word_count = {}
        for word_tuple in words:
            word = word_tuple[0]
            if word in word_count:
                word_count[word] += 1
            else:
                word_count[word] = 1
        
        # 找出重复的单词
        duplicate_words = [word for word, count in word_count.items() if count > 1]
        
        cursor.close()
        db.close()
        
        return duplicate_words
    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return []



## 从主表获取以其外为外键的表

def get_related_tables_by_foreign_key(database, base_table):
    table_structures = chakan_biao_biaojiegou(database)
    if not table_structures:
        return {}

    related_tables = {}
    
    for table, create_stmt in table_structures.items():
        fk_pattern = re.compile(r"FOREIGN KEY \(`(\w+)`\) REFERENCES `(\w+)` \(`(\w+)`\)")
        fks = fk_pattern.findall(create_stmt)
        for fk in fks:
            if fk[1] == base_table:
                related_tables[table] = fk[0]
    
    return related_tables


## 从主表id获取所有以其为外键的表的相应记录

def get_related_records(database, base_table, base_table_id):
    related_records = {}

    db = connect_to_db(database)
    cursor = db.cursor()

    related_tables = get_related_tables_by_foreign_key(database, base_table)

    print(f"查询 {base_table} 表中所有 {base_table_id} 非空的记录，获取已存在于其他表中的相关记录。")
    for table, fk_field in related_tables.items():
        query = f"SELECT * FROM `{table}` WHERE `{fk_field}` = %s"
        cursor.execute(query, (base_table_id,))
        records = cursor.fetchall()
        related_records[table] = records

    cursor.close()
    db.close()
    
    return related_records


def extract_kth_field_from_records(records, k):
    """
    从记录中提取第 k 个字段的值并返回一个列表。
    
    :param records: 记录字典，其中键是表名，值是记录列表。
    :param k: 字段索引，从0开始。
    :return: 包含第 k 个字段值的列表。
    """
    result = []
    for table, record_list in records.items():
        for record in record_list:
            if len(record) > k:
                result.append(record[k])
            else:
                print(f"记录 {record} 中没有第 {k} 个字段")
    return result


##从主键表id获取以其为主键的下一级对象的名称

def get_related_xiayiji_name(database, base_table, base_table_id):

    """
    从cluster的id获取下面的groups名的列表


    从groups的id获取对应的words中的单词列表


    """
    

    related_records = {}

    db = connect_to_db(database)
    cursor = db.cursor()

    related_tables = get_related_tables_by_foreign_key(database, base_table)

    print(f"查询 {base_table} 表中所有 {base_table_id} 非空的记录，获取已存在于其他表中的相关记录。")
    for table, fk_field in related_tables.items():
        query = f"SELECT * FROM `{table}` WHERE `{fk_field}` = %s"
        cursor.execute(query, (base_table_id,))
        records = cursor.fetchall()
        related_records[table] = records

    cursor.close()
    db.close()

    res = extract_kth_field_from_records(related_records, 2)
    
    return res




## 获取words中group_id 字段非空的所有记录的单词

def extract_words_with_non_null_group_id(database):
    """
    遍历 words 表，将其 group_id 字段非空的所有记录存入列表 A，
    从 A 的记录的 word 字段提取单词，将所有这些单词放入列表 LB，并返回 LB。
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 查询 words 表中所有 group_id 非空的记录
        cursor.execute("SELECT * FROM words WHERE group_id IS NOT NULL")
        records = cursor.fetchall()

        # 列表 A 存储所有符合条件的记录
        A = records

        # 从 A 的记录的 word 字段提取单词，将其存入列表 LB
        LB = [record[2] for record in A]  # 假设 word 字段是第三个字段，索引为 2

        cursor.close()
        db.close()

        return LB

    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return []

##提取get_related_records(database, "groups", k)返回值中的单词

def extract_words_from_related_records(related_records):
    """
    从 get_related_records 返回的 related_records 中提取 words 表的记录中的单词。
    
    :param related_records: 由 get_related_records 返回的字典，包含表名和对应的记录
    :return: 包含提取出的单词的列表
    """
    words_list = []
    
    # 获取 words 表的相关记录
    if 'words' in related_records:
        words_records = related_records['words']
        
        # 提取每个记录的 word 字段并存入列表
        for record in words_records:
            word = record[2]  # 假设 word 字段是第三个字段，索引为 2
            words_list.append(word)
    
    return words_list


#删除表指定id记录的函数


def delete_record_by_id(database, table, id_value, id_column="id"):
    """
    从指定数据库和表中删除指定主键值的记录。

    参数:
        database: 数据库名称
        table: 表名称
        id_value: 要删除的记录主键值
        id_column: 主键字段名，默认是 "id"
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        sql = f"DELETE FROM `{table}` WHERE `{id_column}` = %s"
        cursor.execute(sql, (id_value,))
        db.commit()

        if cursor.rowcount > 0:
            print(f"✅ 成功删除 `{table}` 表中 `{id_column}` = {id_value} 的记录")
        else:
            print(f"⚠️ 未找到 `{table}` 表中 `{id_column}` = {id_value} 的记录")

        cursor.close()
        db.close()
    except Exception as e:
        print(f"❌ 删除失败：{e}")



#通用批量删除函数（支持多条件）

def delete_records(database, table, where_field, values, operator="IN"):
    """
    批量删除记录，可按任意字段值批量删除。

    参数:
        database: 数据库名
        table: 表名
        where_field: 条件字段（如 "id", "room_id", "name" 等）
        values: 字段对应的值列表
        operator: SQL 操作符，默认 "IN"，也可以是 "="（用于单值）、"!="、"NOT IN" 等
    """
    if not values:
        print("⚠️ 删除操作被跳过：未提供有效的条件值")
        return

    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 如果是单个值 + "="
        if operator.upper() == "=" and len(values) == 1:
            sql = f"DELETE FROM `{table}` WHERE `{where_field}` = %s"
            cursor.execute(sql, (values[0],))
        else:
            placeholders = ', '.join(['%s'] * len(values))
            sql = f"DELETE FROM `{table}` WHERE `{where_field}` {operator} ({placeholders})"
            cursor.execute(sql, tuple(values))

        db.commit()
        print(f"✅ 删除 `{table}` 表中 `{where_field}` 匹配条件的记录共 {cursor.rowcount} 条")

        cursor.close()
        db.close()

    except Exception as e:
        print(f"❌ 删除失败：{e}")
        
##delete_records("某某住宅", "doors", "id", [1, 2, 3])
##delete_records("某某住宅", "doors", "room_id", [101])
##delete_records("某某住宅", "rooms", "name", ["房间1", "房间2"])


#高级通用删除函数（多字段条件支持）

def delete_records_with_conditions(database, table, conditions: dict):
    """
    删除满足多个字段条件的记录（联合 WHERE 条件）。

    参数:
        database: 数据库名
        table: 表名
        conditions: 字典，键是字段名，值是具体值，如：
                   {"room_id": 5, "door_type": "单开门"}
    """
    if not conditions:
        print("⚠️ 未提供删除条件，操作被跳过。")
        return

    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 构造 WHERE 子句
        where_clauses = []
        values = []

        for field, value in conditions.items():
            where_clauses.append(f"`{field}` = %s")
            values.append(value)

        where_str = " AND ".join(where_clauses)
        sql = f"DELETE FROM `{table}` WHERE {where_str};"

        cursor.execute(sql, tuple(values))
        db.commit()
        print(f"✅ 删除 `{table}` 中符合条件的记录：{cursor.rowcount} 条")

        cursor.close()
        db.close()

    except Exception as e:
        print(f"❌ 删除失败：{e}")

##删除 room_id = 5 且 door_type = "单开门" 的所有门：
##python
##复制
##编辑
##delete_records_with_conditions(
##    database="某某住宅",
##    table="doors",
##    conditions={"room_id": 5, "door_type": "单开门"}
##)
##删除窗宽为 1500 且类型为“推拉窗”的所有窗：
##python
##复制
##编辑
##delete_records_with_conditions(
##    database="某某住宅",
##    table="windows",
##    conditions={"width": 1500, "window_type": "推拉窗"}
##)


#进阶用删除函数（多字段条件支持）


def delete_records_advanced(database, table, conditions):
    """
    高级删除记录函数，根据条件组合删除指定记录。

    参数:
        database (str): 数据库名
        table (str): 表名
        conditions (dict): 复杂条件字典，格式如下：
            {
                "字段名": ("操作符", 值)  # 支持 '=', '!=', 'LIKE', 'IN', 'BETWEEN'
            }

    示例:
        delete_records_advanced("某某住宅", "rooms", {
            "category": ("LIKE", "%厨房%"),
            "depth": ("BETWEEN", (3000, 6000)),
            "id": ("IN", [2, 3, 5])
        })
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        condition_sqls = []
        values = []

        for field, (op, val) in conditions.items():
            if op.upper() == "IN" and isinstance(val, (list, tuple)):
                placeholders = ", ".join(["%s"] * len(val))
                condition_sqls.append(f"`{field}` IN ({placeholders})")
                values.extend(val)
            elif op.upper() == "BETWEEN" and isinstance(val, (list, tuple)) and len(val) == 2:
                condition_sqls.append(f"`{field}` BETWEEN %s AND %s")
                values.extend(val)
            elif op.upper() in ["=", "!=", "LIKE"]:
                condition_sqls.append(f"`{field}` {op} %s")
                values.append(val)
            else:
                print(f"❌ 不支持的操作符或值: {field} {op} {val}")
                return False

        condition_clause = " AND ".join(condition_sqls)
        sql = f"DELETE FROM `{table}` WHERE {condition_clause};"

        cursor.execute(sql, values)
        db.commit()
        print(f"✅ 成功删除记录：{cursor.rowcount} 行")

        cursor.close()
        db.close()
        return True

    except Exception as e:
        print(f"❌ 删除失败：{e}")
        return False
### 删除 id 为 2, 3, 4 且 category 包含“过道”的房间
##delete_records_advanced("某某住宅", "rooms", {
##    "id": ("IN", [2, 3, 4]),
##    "category": ("LIKE", "%过道%")
##})






##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


###            更新表数据



##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&





#### 信息显示

############### 


## 2 查询特定表的第n-m条信息n=start,m=end

def show_records_range(database, table_name, start, end):
    results = []
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 构建查询语句，限制返回从第 start 条到第 end 条记录，使用 LIMIT 和 OFFSET 实现分页查询
        query = f"SELECT * FROM {table_name} ORDER BY id ASC LIMIT %s OFFSET %s"
        cursor.execute(query, (end - start + 1, start - 1))
        
        # 获取从第 start 条到第 end 条记录
        records = cursor.fetchall()
        results.append(records)
        
        # 获取列名
        column_names_query = f"SHOW COLUMNS FROM {table_name}"
        cursor.execute(column_names_query)
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]
        
        # 打印列名
        print(column_names)
        results.append(column_names)
        
        # 打印记录
        for record in records:
            print(record)
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    return results

## 3 针对查询结果提取纯数据

def extract_kth_elements(LS, k):
    # 获取数据部分，跳过列名部分
    data = LS[0]
    
    # 提取第 k 个元素，注意 k 是从 1 开始计数
    result = [record[k - 1] for record in data]
    
    return result






# 4 查询数据库database任何表id号的信息,默认查询数据库database表table_name特定字段的全部记录信息
def show_records_id(database, table_name, record_id=None):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 使用反引号括起表名
        table_name_escaped = f"`{table_name}`"
        
        # 构建查询语句
        if record_id is None:
            query = f"SELECT * FROM {table_name_escaped}"
        else:
            query = f"SELECT * FROM {table_name_escaped} WHERE id = %s"
        
        cursor.execute(query, (record_id,) if record_id is not None else ())
        
        # 获取所有记录
        records = cursor.fetchall()
        
        # 获取列名
        column_names_query = f"SHOW COLUMNS FROM {table_name_escaped}"
        cursor.execute(column_names_query)
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]
        
        # 打印列名
        print(column_names)
        
        # 打印记录
        for record in records:
            print(record)
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

    return records

### 示例用法
##database = "English_wordsstudy"
##table_name = "WordBasicInfo"  # 可以替换为任何表名    
##record_id=17544

##field="phonetic" 


def show_records_id_budayin(database, table_name, record_id=None):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 使用反引号括起表名
        table_name_escaped = f"`{table_name}`"
        
        # 构建查询语句
        if record_id is None:
            query = f"SELECT * FROM {table_name_escaped}"
        else:
            query = f"SELECT * FROM {table_name_escaped} WHERE id = %s"
        
        cursor.execute(query, (record_id,) if record_id is not None else ())
        
        # 获取所有记录
        records = cursor.fetchall()
        
        # 获取列名
        column_names_query = f"SHOW COLUMNS FROM {table_name_escaped}"
        cursor.execute(column_names_query)
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

    return records




##  5 根据表的字段值查询
def show_records_by_field(database, table_name, field_name, field_value):
    results = []
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 构建查询语句
        query = f"SELECT * FROM `{table_name}` WHERE `{field_name}` = %s"
        cursor.execute(query, (field_value,))
        
        # 获取所有记录
        records = cursor.fetchall()
        results.append(records)
        
        # 获取列名
        column_names_query = f"SHOW COLUMNS FROM `{table_name}`"
        cursor.execute(column_names_query)
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]
        
        # 打印列名
        print(column_names)
        results.append(column_names)
        
        # 打印记录
        for record in records:
            print(record)
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    return results

# 6 根据字段列表值批量查询

def batch_show_records_by_field(database, table_name, field_name, field_values_list):
    all_results = []
    for field_value in field_values_list:
        print(f"Records for {field_name} = {field_value}:")
        results = show_records_by_field(database, table_name, field_name, field_value)
        all_results.append(results)
        print("\n" + "-"*50 + "\n")
    return all_results

##database = "English_wordsstudy"
##table_name = "WordBasicInfo"
##field_name = "word"
##LB = ["P.E.", "silent"]
##    
##batch_show_records_by_field(database, table_name, field_name, LB)

##根据指定的被单词包含的字符串查询所有单词并返回 20240629

def find_words_containing_substring(database,substring):
    # 数据库连接参数
    host = "localhost"
    user = "root"
    password = os.getenv("DATABASE_PASSWORD")
        
    # 连接到MySQL数据库
    conn = mysql.connector.connect(
        host=host,
        user=user,
        password=password,
        database=database
    )
    cursor = conn.cursor()
    
    # SQL 查询
    query = "SELECT word FROM words WHERE word LIKE %s"
    
    cursor.execute(query, (f'%{substring}%',))
    
    # 获取结果并存入列表
    words_list = [row[0] for row in cursor.fetchall()]
    
    # 关闭连接
    conn.close()
    
    return words_list






#### 数据更新

###############


##更新数据库指定表特定id的指定字段
def update_field_info(database, table_name, record_id, field, new_value):

    """
    使用show_top_n_records(database, table_name, n)命令可以看到任何表的所有字段，一般至少包括id，外键，一个具体值

    """    
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 构建更新语句
        update_query = f"UPDATE {table_name} SET {field} = %s WHERE id = %s"
        cursor.execute(update_query, (new_value, record_id))
        
        db.commit()
        print(f"Record with id {record_id} in table {table_name} updated successfully. Field '{field}' is now '{new_value}'.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")




##show_top_n_records(database, "wordforms", 10)
##['id', 'word_id', 'word']
##(1, 11700, 'a')
##(2, 11701, 'abandon')
##(3, 11702, 'ability')
##(4, 11703, 'able')
##(5, 11704, 'aboard')
##(6, 11705, 'abolish')
##(7, 11706, 'about')
##(8, 11707, 'above')
##(9, 11708, 'abroad')
##(10, 11709, 'absence')
        
##批量更新记录字段
def batch_update_records(database, table_name, fields, *values_lists):
    if len(fields) != len(values_lists):
        print("Error: The number of fields does not match the number of values lists.")
        return
    
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        for values in zip(*values_lists):
            # 获取当前记录的id
            cursor.execute(f"SELECT id FROM {table_name} WHERE {fields[0]} = %s", (values[0],))
            record_id = cursor.fetchone()
            
            if record_id:
                record_id = record_id[0]
                for field, value in zip(fields, values):
                    update_field_info(database, table_name, record_id, field, value)
            else:
                print(f"No record found with {fields[0]} = {values[0]}")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

###使用方法
##database = "English_wordsstudy"
##table_name = "WordBasicInfo"
##fields = ['word', 'phonetic', 'meaning']
##LB1 = ["P.E.", "silent"]
##LB2 = ["pi:i:-cs", "silent-cs"]
##LB3 = ["n.体育-cs", "a.沉默的-cs"]
##    
##batch_update_records(database, table_name, fields, LB1, LB2, LB3)






### 批量在末端添加记录，删除记录


def append_records(database, table_name, records):
    """
    追加新记录到指定的表中。
    :param database: 数据库名称
    :param table_name: 表名称
    :param records: 要追加的新记录列表，每个记录是一个元组，包含要插入的字段值
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 获取表的列名
        cursor.execute(f"SHOW COLUMNS FROM {table_name}")
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns if column[0] != 'id']  # 排除id列

        # 构建插入SQL语句
        columns_str = ', '.join(column_names)
        placeholders = ', '.join(['%s'] * len(column_names))
        insert_query = f"INSERT INTO {table_name} ({columns_str}) VALUES ({placeholders})"

        # 插入新记录
        cursor.executemany(insert_query, records)
        db.commit()

        print(f"Inserted {cursor.rowcount} records into {table_name} successfully.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

### 示例用法
##database = "Learning_English_vocabulary"
##table_name = "WordBasicInfo"
##
### 假设 WordBasicInfo 表有字段：code, word, phonetic, meaning
##new_records = [
##    (5863, 'newword1', 'newphonetic1', 'new meaning 1'),
##    (5864, 'newword2', 'newphonetic2', 'new meaning 2'),
##    # 添加更多记录
##]
##
##append_records(database, table_name, new_records)

# 将数据库的表，包括最后一条在内的最后n个记录删除

def delete_last_n_records(database, table_name, n):
    """
    删除数据库表中的最后 n 个记录。
    :param database: 数据库名称
    :param table_name: 表名称
    :param n: 要删除的记录数
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()
        
        # 获取表中的记录数
        cursor.execute(f"SELECT COUNT(*) FROM `{table_name}`")
        total_records = cursor.fetchone()[0]

        if n > total_records:
            print(f"Error: Table {table_name} has only {total_records} records, but {n} records were requested for deletion.")
            return

        # 获取最后 n 个记录的 ID
        cursor.execute(f"SELECT id FROM `{table_name}` ORDER BY id DESC LIMIT %s", (n,))
        ids_to_delete = cursor.fetchall()

        # 删除最后 n 个记录
        ids_tuple = tuple(id[0] for id in ids_to_delete)
        if len(ids_tuple) == 1:
            delete_query = f"DELETE FROM `{table_name}` WHERE id = {ids_tuple[0]}"
        else:
            delete_query = f"DELETE FROM `{table_name}` WHERE id IN {ids_tuple}"
        cursor.execute(delete_query)
        db.commit()

        print(f"Deleted the last {n} records from {table_name} successfully.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")



### 示例用法
##database = "Learning_English_vocabulary"
##table_name = "WordBasicInfo"
##n = 5  # 要删除的记录数
##
##delete_last_n_records(database, table_name, n)


##指定id
def append_records_with_id(database, table_name, records):
    """
    追加新记录到指定的表中，并手动指定 ID。
    :param database: 数据库名称
    :param table_name: 表名称
    :param records: 要追加的新记录列表，每个记录是一个元组，包含要插入的字段值
    """
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 获取表的列名
        cursor.execute(f"SHOW COLUMNS FROM `{table_name}`")
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]  # 包括 id 列

        # 构建插入SQL语句
        columns_str = ', '.join(column_names)
        placeholders = ', '.join(['%s'] * len(column_names))
        insert_query = f"INSERT INTO `{table_name}` ({columns_str}) VALUES ({placeholders})"

        # 插入新记录
        cursor.executemany(insert_query, records)
        db.commit()

        print(f"Inserted {cursor.rowcount} records into {table_name} successfully.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

def append_records_with_id(database: str, table_name: str, records: List[Tuple]):
    """
    追加新记录到指定的表中，并手动指定 ID。
    :param database: 数据库名称
    :param table_name: 表名称
    :param records: 要追加的新记录列表，每个记录是一个元组，包含要插入的字段值
    """
    db = None
    cursor = None
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 获取表的列名
        cursor.execute(f"SHOW COLUMNS FROM `{table_name}`")
        columns = cursor.fetchall()
        column_names = [column[0] for column in columns]  # 包括 id 列

        # 构建插入SQL语句
        columns_str = ', '.join([f"`{col}`" for col in column_names])
        placeholders = ', '.join(['%s'] * len(column_names))
        insert_query = f"INSERT INTO `{table_name}` ({columns_str}) VALUES ({placeholders})"

        # 打印生成的插入SQL语句和参数用于调试
        print(insert_query)
        print(records)

        # 插入新记录
        cursor.executemany(insert_query, records)
        db.commit()

        print(f"Inserted {cursor.rowcount} records into {table_name} successfully.")
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    finally:
        if cursor is not None:
            cursor.close()
        if db is not None:
            db.close()


### 确认 projects 表中的记录
##show_records_id(database, "projects")
##
### 插入 clusters 表的记录，手动指定 ID
##records = [
##    (1, 1, "第1集",""), (2, 1, "第2集",""), (3, 1, "第3集",""),
##    (4, 1, "第4集",""), (5, 1, "第5集",""), (6, 1, "第6集",""),
##    (7, 1, "第7集",""), (8, 1, "第8集",""), (9, 1, "第9集","")
##]
##append_records_with_id(database, "clusters", records)



## 更新 groups 表记录的 cluster_id 字段
def update_cluster_id_for_group_names(database, target_cluster_id, names_to_check):
    """
    更新 groups 表中符合条件的记录的 cluster_id 字段。
    
    :param database: 数据库名称
    :param target_cluster_id: 要设置的 cluster_id 值
    :param names_to_check: 单词列表，不区分大小写
    :param case_sensitive_exceptions: 需要区分大小写的单词列表
    """
    case_sensitive_exceptions = ['March', 'May'] 

    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 获取 groups 表中所有的记录
        cursor.execute("SELECT id, name FROM `groups`")
        records = cursor.fetchall()

        # 更新符合条件的记录的 cluster_id
        for record in records:
            group_id, name = record
            # 检查是否在 names_to_check 中
            if name.lower() in [word.lower() for word in names_to_check]:
                # 检查是否为 case_sensitive_exceptions
                if name in case_sensitive_exceptions:
                    # 区分大小写
                    if name in names_to_check:
                        cursor.execute("UPDATE `groups` SET cluster_id = %s WHERE id = %s", (target_cluster_id, group_id))
                else:
                    # 不区分大小写
                    cursor.execute("UPDATE `groups` SET cluster_id = %s WHERE id = %s", (target_cluster_id, group_id))

        db.commit()

        print("Successfully updated the cluster_id for matching records in the groups table.")
        
        cursor.close()
        db.close()
    except mysql.connector.Error as err:
        print(f"Error: {err}")

### 示例用法
##names_to_check = ['rat', 'questionnaire', 'satisfy', 'software', 'shout', 'give', 'kilo']
##
##case_sensitive_exceptions = ['march', 'may']  # 假设这两个单词区分大小写
##
##update_cluster_id_for_group_names(database, 1, names_to_check, case_sensitive_exceptions)



##获取特定后缀文件的包含完整路径后缀名的文件名
def get_all_files_with_extensions(directory, extensions):
    files_with_extensions = []
    
    # Ensure extensions are in lowercase and start with a dot
    extensions = [ext.lower() if ext.startswith('.') else '.' + ext.lower() for ext in extensions]
    
    # Walk through the directory structure
    for root, _, files in os.walk(directory):
        for file in files:
            if any(file.lower().endswith(ext) for ext in extensions):
                # Construct the full file path
                file_path = os.path.join(root, file)
                files_with_extensions.append(file_path)
    
    return files_with_extensions

##directory = '/path/to/your/folder'  # Replace with the path to your folder
##extensions = ['.png', '.dwg']  # Replace with the desired file extensions


##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


###         整数据库迁移



##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&


## 数据库导出数据

def export_database_pyfunc(database_name):

    output_file =os.path.join( 'C:/Users/Administrator/',database_name + '.sql')
        
    res = export_database(database_name, output_file)

    if res == 0:

        print("数据库导出数据失败！")

    elif res == 1:

        print("数据库导出数据成功！")

    else:

        pass

    return res





## 用函数创建数据库


def create_database_pyfunc(filename_chunming,new_database_name):

    """
    实测输入文件路径不能随便放

    r'C:/Users/Administrator/0621wanEnglish_wordsstudy_backup_utf8.sql'

    """
    input_file = os.path.join(r'C:/Users/Administrator/',filename_chunming+'.sql')   

    delete_database(new_database_name)

    create_database_if_not_exists(new_database_name)

    import_database(new_database_name, input_file)
    
    time.sleep(5)

    connect_to_db_no_db()

    res = login_and_check_database(new_database_name)

    if res == 0:

        print("数据库创建失败！")

    elif res == 1:

        print("数据库创建成功！")

    else:

        pass

    return res


##&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&




## 图

db_config = {
    'host': "localhost",
    'user': "root",
    'password': os.getenv("DATABASE_PASSWORD"),
    'database': "CADdata"
}


def insert_vertices(vertices_data, database="CADdata"):
    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS vertices (
                vertex_id INT PRIMARY KEY,
                id INT,
                x FLOAT,
                y FLOAT,
                z FLOAT,
                additional_field VARCHAR(255)
            )
        """)
        cursor.executemany("INSERT INTO vertices (vertex_id, id, x, y, z, additional_field) VALUES (%s, %s, %s, %s, %s, %s)", vertices_data)
        conn.commit()
        print("顶点数据插入成功！")
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    finally:
        cursor.close()
        conn.close()


def insert_edges(edges_data, database="CADdata"):
    try:
        conn = connect_to_db(database)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS edges (
                edge_id INT PRIMARY KEY,
                id INT,
                start_vertex_id INT,
                end_vertex_id INT,
                length FLOAT,
                additional_field VARCHAR(255),
                FOREIGN KEY (start_vertex_id) REFERENCES vertices(vertex_id),
                FOREIGN KEY (end_vertex_id) REFERENCES vertices(vertex_id)
            )
        """)
        cursor.executemany("INSERT INTO edges (edge_id, id, start_vertex_id, end_vertex_id, length, additional_field) VALUES (%s, %s, %s, %s, %s, %s)", edges_data)
        conn.commit()
        print("边数据插入成功！")
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    finally:
        cursor.close()
        conn.close()





##重置ID自动递增

def reset_auto_increment(database, table_name, new_value):
    try:
        db = connect_to_db(database)
        cursor = db.cursor()

        # 构建SQL语句
        sql = f"ALTER TABLE {table_name} AUTO_INCREMENT = {new_value};"

        # 执行SQL语句
        cursor.execute(sql)

        # 提交更改
        db.commit()

        print(f"Auto-increment value for table {table_name} has been reset to {new_value}.")
    
    except mysql.connector.Error as err:
        print(f"Error: {err}")
    
    finally:
        # 关闭数据库连接
        cursor.close()
        db.close()



##导出为Excel   show_records_id(database, "groups")


def export_mysql_to_excel(user="root", host="localhost", database="Learning_English_vocabulary",
                          output_path=r"C:\Users\Administrator\Learning_English_vocabulary导出数据.xlsx"):
    """
    将 MySQL 数据库中所有表导出为一个 Excel 文件，每张表一个 sheet。
    """
    password = os.getenv("DATABASE_PASSWORD")

    try:
        engine = create_engine(f"mysql+pymysql://{user}:{password}@{host}/{database}?charset=utf8mb4")

        tables = pd.read_sql("SHOW TABLES", engine)
        table_names = tables.iloc[:, 0].tolist()

        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            for table in table_names:
                safe_table_name = f"`{table}`"  # 包裹表名避免关键字冲突
                df = pd.read_sql(f"SELECT * FROM {safe_table_name}", engine)
                df.to_excel(writer, sheet_name=table[:31], index=False)

        print(f"✅ 已成功将数据库 '{database}' 中所有表导出到 Excel：{output_path}")
        return True

    except Exception as e:
        print(f"❌ 导出失败：{e}")
        return False


def export_mysql_schema_to_excel(user='root',host='localhost', database='your_database_name',
                                  output_path=r'C:\Users\Administrator\Learning_English_vocabulary导出结构.xlsx'):#导出数据库结构信息
    """
    将 MySQL 数据库的表结构信息导出到 Excel 文件中（不含数据）。
    
    """
    password = os.getenv("DATABASE_PASSWORD")
    
    try:
        conn = pymysql.connect(host=host, user=user, password=password, database=database)
        cursor = conn.cursor()

        cursor.execute("SHOW TABLES")
        tables = [row[0] for row in cursor.fetchall()]

        writer = pd.ExcelWriter(output_path, engine='openpyxl')

        for table in tables:
            cursor.execute(f"SHOW FULL COLUMNS FROM `{table}`")
            columns = cursor.fetchall()
            desc = [col[0] for col in cursor.description]  # 字段名
            df = pd.DataFrame(columns, columns=desc)
            df.to_excel(writer, sheet_name=table[:31], index=False)
            print(f"✅ 表结构已导出: {table}（{len(columns)}列）")

        writer.close()
        print(f"📁 所有表结构已保存至: {output_path}")
        return True
    except Exception as e:
        print(f"❌ 导出表结构失败：{e}")
        return False





#从Excel恢复



def restore_database_from_excel(excel_path, database, user="root", host="localhost", table_exists="replace"):
    """
    将 Excel 文件中的所有 sheet 作为表重建到指定 MySQL 数据库中。

    参数:
        excel_path (str): Excel 文件路径
        database (str): 目标数据库名（需已存在）
        user (str): MySQL 用户名
        host (str): 数据库地址
        table_exists (str): 如果表已存在，"replace" 表示覆盖，"append" 表示追加

    返回:
        True 成功恢复，False 失败
    """
    password = os.getenv("DATABASE_PASSWORD")

    try:
        # 建立数据库连接
        engine = create_engine(f"mysql+pymysql://{user}:{password}@{host}/{database}?charset=utf8mb4")

        # 读取 Excel 所有工作表
        excel_data = pd.read_excel(excel_path, sheet_name=None, engine="openpyxl")

        for sheet_name, df in excel_data.items():
            # 规范化表名：去除空格和特殊字符
            safe_table_name = sheet_name.strip().replace(" ", "_").replace("-", "_")
            df.to_sql(safe_table_name, engine, index=False, if_exists=table_exists)
            print(f"✅ 已写入表: {safe_table_name}，行数: {len(df)}")

        print(f"🎉 数据恢复完成，Excel 文件：{excel_path} → 数据库：{database}")
        return True

    except Exception as e:
        print(f"❌ 数据恢复失败：{e}")
        return False



#完全恢复

def restore_database_from_schema_and_data(schema_path, data_path, database_name,
                                          user='root', host='localhost', password=None):
    if password is None:
        password = os.getenv("DATABASE_PASSWORD")
    try:
        conn = pymysql.connect(host=host, user=user, password=password)
        cursor = conn.cursor()
        cursor.execute(f"CREATE DATABASE IF NOT EXISTS `{database_name}` DEFAULT CHARACTER SET utf8mb4")
        print(f"✅ 已创建新数据库：{database_name}")
        conn.select_db(database_name)

        schema_excel = pd.ExcelFile(schema_path)

        for table in schema_excel.sheet_names:
            df = pd.read_excel(schema_excel, sheet_name=table)
            cols = []
            primary_keys = []

            for _, row in df.iterrows():
                field = row['Field']
                col_type = row['Type']
                nullable = "NULL" if row['Null'] == "YES" else "NOT NULL"
                default = row['Default']
                default_clause = f"DEFAULT '{default}'" if pd.notna(default) and default != 'NULL' else ""
                extra = row['Extra'] if pd.notna(row['Extra']) else ""

                col_def = f"`{field}` {col_type} {nullable} {default_clause} {extra}".strip()
                cols.append(col_def)

                if 'auto_increment' in extra.lower():
                    primary_keys.append(field)

            if primary_keys:
                pk_clause = f", PRIMARY KEY (`{primary_keys[0]}`)"
            else:
                pk_clause = ""

            full_stmt = f"CREATE TABLE `{table}` ({', '.join(cols)}{pk_clause}) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4"
            cursor.execute(full_stmt)
            print(f"✅ 已创建表结构: {table}")

        engine = create_engine(f"mysql+pymysql://{user}:{password}@{host}/{database_name}?charset=utf8mb4")
        data_excel = pd.ExcelFile(data_path)

        for table in data_excel.sheet_names:
            df = pd.read_excel(data_excel, sheet_name=table)
            df.to_sql(table, engine, if_exists='append', index=False)
            print(f"✅ 已导入数据表: {table}，行数: {len(df)}")

        return True
    except Exception as e:
        print(f"❌ 恢复失败：{e}")
        return False



import pandas as pd
from sqlalchemy import create_engine




from sqlalchemy import create_engine, text
import pandas as pd
from typing import Iterable, Dict, List, Tuple, Any, Union

# 1) 连接：加上 charset，避免中文释义乱码
def connect_to_db_x():
    engine = create_engine(
        "mysql+mysqlconnector://root:sy-UXgdZ7yhjH3Fa1mYBaSA598fA399444f9a6fFcDcC3Ed6925@localhost/Learning_English_vocabulary?charset=utf8mb4"
    )
    return engine

# 2) 按组取 word（使用参数化，避免拼接 SQL）
def get_words_by_group(group_id: int) -> pd.DataFrame:
    engine = connect_to_db_x()
    sql = text("""
        SELECT word
        FROM `words`
        WHERE group_id = :gid
        ORDER BY id
        LIMIT 100
    """)
    with engine.connect() as conn:
        df = pd.read_sql(sql, conn, params={"gid": group_id})
    return df

# 3) 打印并返回 DataFrame（关键：一定要 return）
def find_words_in_rat_group(k: int = 1) -> pd.DataFrame:
    df = get_words_by_group(k)

    # 仅为了展示完整列表（非必须）
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)

    print(f"{k} 组中的单词：")
    print(df)
    return df  # ★★★ 一定要返回

# ================== 以下保持你之前的合并逻辑不变 ==================
try:
    from sqlalchemy.engine import Engine, Connection
    from sqlalchemy import text as _text
except Exception:
    Engine = Connection = tuple()
    _text = None

try:
    from mysql.connector.connection_cext import CMySQLConnection as MySQLConn
except Exception:
    MySQLConn = tuple()

def _fetch_words_from_table(
    conn: Union["Engine","Connection","MySQLConn"],
    table: str,
    words: List[str],
) -> List[Dict[str, Any]]:
    if not words:
        return []
    cols = "word, phonetic, meaning"
    if isinstance(conn, MySQLConn):
        ph = ", ".join(["%s"] * len(words))
        sql = f"SELECT {cols} FROM `{table}` WHERE word IN ({ph})"
        cur = conn.cursor(dictionary=True)
        cur.execute(sql, words)
        rows = cur.fetchall()
        cur.close()
        return rows
    elif isinstance(conn, (Engine, Connection)) and _text is not None:
        placeholders = ", ".join([f":w{i}" for i in range(len(words))])
        sql = _text(f"SELECT {cols} FROM `{table}` WHERE word IN ({placeholders})")
        params = {f"w{i}": w for i, w in enumerate(words)}
        if isinstance(conn, Engine):
            with conn.connect() as c:
                res = c.execute(sql, params)
                return [dict(r._mapping) for r in res]
        else:
            res = conn.execute(sql, params)
            return [dict(r._mapping) for r in res]
    else:
        raise TypeError("不支持的数据库连接类型")

def get_group_words_info(
    k: int,
    conn: Union["Engine","Connection","MySQLConn"],
    *,
    table_priority: Tuple[str, ...] = ("wordbasicinfo", "words"),
    keep_order_from_source: bool = True,
) -> pd.DataFrame:
    df_words = find_words_in_rat_group(k=k)   # 现在能返回 DataFrame 了
    if "word" not in df_words.columns:
        raise ValueError("find_words_in_rat_group(k) 返回不含 'word' 列")

    words_raw: List[str] = df_words["word"].astype(str).tolist()

    # 去重保序
    seen = set()
    words_unique: List[str] = []
    for w in words_raw:
        if w not in seen:
            seen.add(w)
            words_unique.append(w)

    # 按优先表取数并合并
    collected: Dict[str, Dict[str, Any]] = {}
    for tb in table_priority:
        rows = _fetch_words_from_table(conn, tb, words_unique)
        for r in rows:
            w = r.get("word")
            if w and w not in collected:
                collected[w] = {
                    "word": w,
                    "phonetic": r.get("phonetic"),
                    "meaning": r.get("meaning"),
                    "source_table": tb,
                }

    # 组装输出（按原顺序）
    order_iter: Iterable[str] = words_raw if keep_order_from_source else sorted(set(words_raw))
    out_rows: List[Dict[str, Any]] = []
    for w in order_iter:
        info = collected.get(w)
        if info is None:
            out_rows.append({"word": w, "phonetic": None, "meaning": None, "source_table": None})
        else:
            out_rows.append(info)

    return pd.DataFrame(out_rows, columns=["word", "phonetic", "meaning", "source_table"])

def format_word_rows(df: pd.DataFrame) -> List[str]:
    lines = []
    for _, r in df.iterrows():
        w = str(r.get("word", "") or "")
        ph = str(r.get("phonetic", "") or "")
        me = str(r.get("meaning", "") or "")
        lines.append(f"{w}\t[{ph}]\t{me}" if ph else f"{w}\t{me}")
    return lines


def get_words_info_by_list(words: list[str], conn, table_priority=("wordbasicinfo","words")):
    # 去重保序
    seen, uniq = set(), []
    for w in words:
        if w not in seen:
            seen.add(w); uniq.append(w)

    # 复用你已有的 _fetch_words_from_table
    collected = {}
    for tb in table_priority:
        rows = _fetch_words_from_table(conn, tb, uniq)
        for r in rows:
            w = r.get("word")
            if w and w not in collected:
                collected[w] = {
                    "word": w,
                    "phonetic": r.get("phonetic"),
                    "meaning": r.get("meaning"),
                    "source_table": tb,
                }

    # 输出按原词序
    out = []
    for w in words:
        out.append(collected.get(w, {"word": w, "phonetic": None, "meaning": None, "source_table": None}))
    return pd.DataFrame(out, columns=["word","phonetic","meaning","source_table"])

#导出为excel/word

def export_group(k: int, engine, *, to_excel: str|None=None, to_docx: str|None=None):
    df_info = get_group_words_info(k=k, conn=engine)  # 走方案A现有函数
    if to_excel:
        df_info.to_excel(to_excel, index=False)
    if to_docx:
        from docx import Document
        doc = Document()
        for line in format_word_rows(df_info):
            doc.add_paragraph(line)
        doc.save(to_docx)
    return df_info










